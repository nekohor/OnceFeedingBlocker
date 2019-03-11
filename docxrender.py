import docx
from docx.shared import Inches
import diagnosis
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches


class DOCXRender:

    def __init__(self, ctx, rec, trnd):
        self.ctx = ctx
        self.rec = rec
        self.trnd = trnd

        self.level1 = 1
        self.level2 = 3
        self.level3 = 4

    def build_docx(self):
        doc = docx.Document("base/docx_base.docx")
        doc.add_heading("重点产品高牌号硅钢一次投料合格率部分", self.level1)
        doc.add_heading("实际汇总", self.level2)

        self.ctx.add_one_table(doc, self.trnd.get_table_for_doc())

        doc.add_heading("趋势与分析", self.level2)

        tag_list_map = {
            "weekly": ["main"],
            "monthly": self.ctx.tags
        }
        for tag in tag_list_map[self.ctx.frequency]:
            item_name = self.trnd.get_trend_item_name(tag)
            doc.add_heading(item_name, self.level3)

            doc.add_paragraph(
                "    {}命中情况的趋势图如下图所示".format(item_name))

            doc.add_picture(
                self.rec.get_item_file_name(
                    "{}_trend.png".format(tag)
                ),
                height=Inches(3.2)
            )

            if tag == "main" and self.ctx.frequency == "weekly":
                diagnosis.week_diagnosis(doc)
            else:
                self.add_pics(doc, tag)

        doc.save(
            self.rec.get_item_file_name(
                '{}{}年{}生产质量分析.docx'.format(
                    item_name,
                    self.ctx.current_year,
                    self.trnd.get_cur_col()
                )
            )
        )

    def add_pics(self, doc, tag):

        cfg_table = self.ctx.plot_config
        for idx in cfg_table.index:
            cfg = cfg_table.loc[idx]

            plot_title = cfg["title"]
            item_name = self.trnd.get_trend_item_name(tag)

            doc.add_paragraph(
                "    {}{}情况如下图所示。".format(
                    item_name, plot_title
                )
            )

            plot_file_name = (
                self.rec.get_item_file_name(
                    "{}_{}.png".format(tag, plot_title)
                )
            )
            doc.add_picture(plot_file_name, height=Inches(3.2))
